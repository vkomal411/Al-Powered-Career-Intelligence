import re
import io
import logging
from typing import Optional, List

import pdfplumber
import docx

logger = logging.getLogger("resume_parser")

from app.ats.text_cleaner import clean_resume_text
from app.ats.contact_detector import (
    detect_contact_info,
    contact_score,
)
from app.ats.section_detector import (
    detect_sections,
    section_score,
)
from app.ats.skill_detector import (
    detect_skills,
    skill_score,
)
from app.ats.score import calculate_ats_score
from app.ats.suggestions import generate_suggestions


# ------------------------------------------------------------------
# Skill keywords (used as fallback if needed)
# ------------------------------------------------------------------

SKILL_KEYWORDS = [
    "python", "java", "javascript", "typescript", "c++", "c#", "c", "sql", "nosql",
    "react", "next.js", "angular", "vue", "node.js", "express", "django", "fastapi",
    "flask", "spring boot", "spring", ".net", "ruby", "ruby on rails", "go", "golang",
    "rust", "kotlin", "swift", "dart", "flutter", "react native",
    "postgresql", "mysql", "mongodb", "redis", "sqlite", "oracle", "firebase",
    "docker", "kubernetes", "aws", "azure", "gcp", "heroku", "vercel", "netlify",
    "git", "github", "gitlab", "bitbucket", "linux", "unix", "bash", "shell",
    "machine learning", "deep learning", "nlp", "computer vision", "data science",
    "pandas", "numpy", "tensorflow", "pytorch", "scikit-learn", "keras", "opencv",
    "matplotlib", "seaborn", "scipy",
    "html", "css", "sass", "tailwind", "bootstrap", "material ui",
    "rest api", "graphql", "websocket", "microservices", "ci/cd", "devops",
    "agile", "scrum", "jira", "confluence",
    "data analysis", "data visualization", "excel", "power bi", "tableau",
    "selenium", "cypress", "jest", "mocha", "junit",
    "wordpress", "shopify", "figma", "photoshop", "canva",
    "blockchain", "solidity", "web3",
    "r", "matlab", "sas", "spss", "hadoop", "spark", "kafka", "airflow",
    "elastic search", "nginx", "apache", "jenkins", "terraform", "ansible",
]

EMAIL_REGEX = re.compile(
    r"[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+",
)

PHONE_REGEX = re.compile(
    r"(?:\+?\d{1,3}[-.\s]?)?"            # optional country code
    r"(?:\(?\d{2,5}\)?[-.\s]?)?"          # optional area code
    r"\d{3,5}[-.\s]?\d{3,5}"             # main number
)


# ------------------------------------------------------------------
# Text Extraction
# ------------------------------------------------------------------

def extract_text_from_pdf(file_bytes: bytes) -> str:
    text_chunks = []
    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_chunks.append(page_text)
    except Exception as e:
        raise ValueError(f"Could not read PDF file: {str(e)}")

    extracted = "\n".join(text_chunks).strip()
    if not extracted:
        raise ValueError("No readable text found in this PDF. If it is a scanned document or image, please convert it to text format.")
    return extracted


def extract_text_from_docx(file_bytes: bytes) -> str:
    try:
        document = docx.Document(io.BytesIO(file_bytes))
    except Exception as e:
        raise ValueError(f"Invalid or corrupted Word (.docx) file: {str(e)}")

    full_text = []

    # Extract paragraphs
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            full_text.append(paragraph.text.strip())

    # Extract tables (many resumes use tables for education/experience)
    for table in document.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                # deduplicate adjacent identical cell text (e.g. merged cells)
                unique_row_cells = []
                for t in row_text:
                    if not unique_row_cells or unique_row_cells[-1] != t:
                        unique_row_cells.append(t)
                full_text.append(" — ".join(unique_row_cells))

    extracted = "\n".join(full_text).strip()
    if not extracted:
        raise ValueError("No text content found inside this Word document.")
    return extracted


def extract_text(filename: str, file_bytes: bytes) -> str:
    lower = filename.lower()

    if lower.endswith(".pdf"):
        return extract_text_from_pdf(file_bytes)

    elif lower.endswith(".docx"):
        return extract_text_from_docx(file_bytes)

    raise ValueError(
        "Unsupported file type. Please upload a PDF or DOCX file."
    )


# ------------------------------------------------------------------
# Contact Extraction
# ------------------------------------------------------------------

def extract_email(text: str) -> Optional[str]:
    match = EMAIL_REGEX.search(text)
    return match.group(0) if match else None


def extract_phone(text: str) -> Optional[str]:
    match = PHONE_REGEX.search(text)
    return match.group(0).strip() if match else None


# ------------------------------------------------------------------
# Legacy Skill Extraction (fallback)
# ------------------------------------------------------------------

def extract_skills(text: str) -> List[str]:
    text_lower = text.lower()

    found = [
        skill
        for skill in SKILL_KEYWORDS
        if re.search(r'\b' + re.escape(skill) + r'\b', text_lower)
    ]

    return sorted(set(found))


# ------------------------------------------------------------------
# Section Isolator Helper
# ------------------------------------------------------------------

SECTION_HEADERS = [
    "education", "academic", "qualifications",
    "experience", "work history", "employment", "internship",
    "projects", "academic projects", "personal projects",
    "skills", "technical skills", "core skills", "core competencies",
    "certifications", "certificates", "licenses",
    "summary", "profile", "objective",
    "achievements", "awards", "honors",
    "hobbies", "interests", "activities",
    "declaration", "references", "personal details", "personal information",
]


def _isolate_section(text: str, section_names: List[str], exit_names: List[str]) -> List[str]:
    """Extract lines belonging to a specific section of the resume."""
    lines = text.splitlines()
    in_section = False
    section_lines: List[str] = []

    for line in lines:
        stripped = line.strip()
        l_lower = stripped.lower()
        if not l_lower:
            continue

        # Check if this is a section header
        is_any_header = any(h in l_lower for h in SECTION_HEADERS) and len(l_lower) < 45

        if any(s in l_lower for s in section_names) and len(l_lower) < 45:
            in_section = True
            continue

        if in_section:
            # Exit when we hit a different section header
            if is_any_header and any(e in l_lower for e in exit_names):
                break
            # Also exit on any unrelated section header
            if is_any_header and not any(s in l_lower for s in section_names):
                break
            section_lines.append(stripped)

    return section_lines


# ------------------------------------------------------------------
# Comprehensive Extractors
# ------------------------------------------------------------------

def extract_name(text: str) -> Optional[str]:
    # 1. Try spaCy NER for PERSON entity extraction
    try:
        from app.ai.spacy_parser import parse_resume_with_spacy
        spacy_result = parse_resume_with_spacy(text[:1000])  # Scan top header region
        for entity in spacy_result.get("entities", []):
            if entity.get("label") == "PERSON":
                candidate_name = entity.get("text", "").strip()
                cleaned_name = re.sub(r"[^a-zA-Z\s.]", "", candidate_name).strip()
                words = cleaned_name.split()
                if 1 <= len(words) <= 4 and all(w.replace(".", "").isalpha() for w in words):
                    return cleaned_name.title()
    except Exception as exc:
        logger.debug("spaCy name extraction skipped or unavailable (%s), falling back to rule-based parser", exc)

    # 2. Rule-based line parsing fallback (Authoritative for non-standard layouts)
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    skip_kw = [
        "http", "github.com", "linkedin.com", "resume", "curriculum",
        "page", "objective", "summary", "profile", "address",
        "phone", "mobile", "email",
    ]
    for line in lines[:8]:
        if EMAIL_REGEX.search(line) or PHONE_REGEX.search(line):
            continue
        if any(kw in line.lower() for kw in skip_kw):
            continue
        cleaned = re.sub(r"[^a-zA-Z\s.]", "", line).strip()
        words = cleaned.split()
        if 1 <= len(words) <= 5 and all(len(w) >= 1 and w.replace(".", "").isalpha() for w in words):
            if cleaned.lower() in [h for h in SECTION_HEADERS]:
                continue
            return cleaned.title()
    return None



# ── Education Extractor ──

DEGREE_KEYWORDS = [
    "bachelor", "master", "phd", "ph.d", "b.tech", "m.tech", "b.e", "m.e",
    "b.s", "m.s", "b.sc", "m.sc", "bba", "mba", "bca", "mca", "associate",
    "diploma", "doctorate", "b.a", "m.a", "b.com", "m.com",
    "intermediate", "high school", "secondary school", "hsc", "sslc",
]

INSTITUTION_KEYWORDS = [
    "university", "college", "institute", "school", "academy", "polytechnic",
    "campus", "iit", "nit", "bits",
]

EXCLUDE_EDUCATION_KEYWORDS = [
    "date of birth", "dob", "father", "mother", "gender", "sex", "marital",
    "nationality", "language", "address", "declaration", "passport", "hobby",
    "hobbies", "reference", "personal details", "profile summary",
    "phone", "mobile", "email", "linkedin", "github", "objective",
]


def clean_education_text(val: str) -> str:
    # 1. Remove leading bullets, numbers, dashes
    val = re.sub(r"^[•\-●○▪*\d.\s]+", "", val)

    # 2. Remove DOB patterns
    val = re.sub(r"\(?\s*date\s+of\s+birth\s*:?\s*[\d/.-]+\s*\)?", "", val, flags=re.IGNORECASE)
    val = re.sub(r"\(?\s*dob\s*:?\s*[\d/.-]+\s*\)?", "", val, flags=re.IGNORECASE)
    val = re.sub(r"\bdate\s+of\s+birth\b[^\n—,]*", "", val, flags=re.IGNORECASE)

    # 3. Remove Expected Graduation / Graduation Year / Passed Out
    val = re.sub(r"\(?\s*expected\s+graduation\s*:?\s*\d{0,4}\s*\)?", "", val, flags=re.IGNORECASE)
    val = re.sub(r"\(?\s*graduation\s+year\s*:?\s*\d{0,4}\s*\)?", "", val, flags=re.IGNORECASE)
    val = re.sub(r"\(?\s*passed\s+out\s*:?\s*\d{0,4}\s*\)?", "", val, flags=re.IGNORECASE)

    # 4. Remove CGPA / Percentage / Marks / Grade scores
    val = re.sub(r"\(?\s*(?:cgpa|gpa|percentage|percentage%|marks|grade)\s*:?\s*[\d.]+%?\s*\)?", "", val, flags=re.IGNORECASE)
    val = re.sub(r"\b\d{1,2}(?:\.\d+)?%\b", "", val)  # e.g. 93.3% or 85%

    # 5. Clean up redundant spaces and empty dashes
    val = re.sub(r"\s+", " ", val)
    val = re.sub(r"\s*—\s*—\s*", " — ", val)
    val = re.sub(r"^\s*[:,\-—]\s*", "", val)
    val = re.sub(r"\s*[:,\-—]\s*$", "", val)
    return val.strip()


def _is_education_relevant(line: str) -> bool:
    l = line.lower()
    has_degree = any(re.search(r'\b' + re.escape(kw) + r'\b', l) for kw in DEGREE_KEYWORDS)
    has_institution = any(kw in l for kw in INSTITUTION_KEYWORDS)
    return has_degree or has_institution


def classify_education_rank(entry_str: str) -> int:
    l = entry_str.lower()

    # Rank 1: Higher Education / Degree (UG / PG / PhD)
    higher_tokens = [
        "bachelor", "master", "phd", "ph.d", "b.tech", "m.tech",
        "b.e", "m.e", "b.s", "m.s", "b.sc", "m.sc", "bba", "mba",
        "bca", "mca", "doctorate", "b.a", "m.a", "b.com", "m.com",
        "diploma", "associate", "university",
    ]
    if any(re.search(r'\b' + re.escape(dt) + r'\b', l) for dt in higher_tokens):
        return 1

    # Rank 2: Secondary Education (Intermediate / 12th / Junior College)
    secondary_tokens = [
        "intermediate", "junior college", "high school", "secondary",
        "hsc", "12th", "class xii", "plus two", "+2", "higher secondary",
    ]
    if any(st in l for st in secondary_tokens):
        return 2

    # Rank 3: Primary / 10th / SSC
    primary_tokens = [
        "primary", "elementary", "sslc", "10th", "class x",
        "school", "matriculation", "ssc",
    ]
    if any(pt in l for pt in primary_tokens):
        return 3

    return 4


def extract_education_history(text: str) -> List[dict]:
    exit_names = [
        "experience", "work history", "projects", "skills",
        "certifications", "summary", "hobbies", "declaration", "references", "personal",
    ]
    section_lines = _isolate_section(
        text,
        section_names=["education", "academic qualification", "academic background", "qualifications"],
        exit_names=exit_names,
    )

    source_lines = section_lines if section_lines else text.splitlines()

    valid_lines: List[str] = []
    for raw_line in source_lines:
        line_str = raw_line.strip()
        if not line_str or len(line_str) < 3:
            continue

        cleaned = clean_education_text(line_str)
        if not cleaned or len(cleaned) < 3:
            continue

        cl = cleaned.lower()

        if any(ex in cl for ex in EXCLUDE_EDUCATION_KEYWORDS):
            continue
        if EMAIL_REGEX.search(cleaned) or PHONE_REGEX.search(cleaned):
            continue

        if _is_education_relevant(cleaned):
            valid_lines.append(cleaned)

    grouped: List[List[str]] = []
    current: List[str] = []

    for line_str in valid_lines:
        cl = line_str.lower()
        has_degree = any(re.search(r'\b' + re.escape(kw) + r'\b', cl) for kw in DEGREE_KEYWORDS)
        has_institution = any(kw in cl for kw in INSTITUTION_KEYWORDS)

        if current and (has_degree or (has_institution and not has_degree)):
            current_text = " ".join(current).lower()
            current_has_degree = any(re.search(r'\b' + re.escape(kw) + r'\b', current_text) for kw in DEGREE_KEYWORDS)

            if current_has_degree and (has_degree or has_institution):
                grouped.append(current)
                current = [line_str]
                continue

        current.append(line_str)

    if current:
        grouped.append(current)

    entries: List[str] = []
    for group in grouped:
        merged = " — ".join(group)
        merged = clean_education_text(merged)
        if merged and len(merged) >= 5:
            entries.append(merged)

    entries.sort(key=classify_education_rank)

    seen: set = set()
    edu_list: List[dict] = []
    for entry in entries:
        key = entry.lower().strip()
        if key not in seen:
            seen.add(key)
            edu_list.append({"degree_or_institution": entry[:250]})
        if len(edu_list) >= 5:
            break

    return edu_list


# ── Experience Extractor ──

JOB_TITLE_KEYWORDS = [
    "software engineer", "developer", "web developer", "full stack",
    "frontend", "backend", "devops", "data scientist", "data analyst",
    "data engineer", "machine learning", "ai engineer",
    "manager", "team lead", "tech lead", "architect", "director",
    "intern", "trainee", "apprentice", "fresher",
    "analyst", "consultant", "administrator", "coordinator",
    "designer", "ui/ux", "product manager", "project manager",
    "qa engineer", "test engineer", "tester", "sqa",
    "specialist", "associate", "executive", "officer",
    "research assistant", "teaching assistant",
]


def extract_experience_history(text: str) -> List[dict]:
    # Try section-isolated extraction first
    exit_names = [
        "education", "projects", "skills", "certifications",
        "summary", "hobbies", "declaration", "references", "personal",
    ]
    section_lines = _isolate_section(
        text,
        section_names=["experience", "work history", "employment", "internship"],
        exit_names=exit_names,
    )

    source = section_lines if section_lines else text.splitlines()
    exp_list: List[dict] = []
    seen: set = set()

    for raw_line in source:
        line_str = raw_line.strip()
        if not line_str or len(line_str) < 5:
            continue

        l_lower = line_str.lower()

        # Skip noise lines
        if EMAIL_REGEX.search(line_str) or PHONE_REGEX.search(line_str):
            continue
        if any(ex in l_lower for ex in ["date of birth", "dob", "gender", "nationality"]):
            continue

        # Check for job-title match
        has_title = any(kw in l_lower for kw in JOB_TITLE_KEYWORDS)
        # Also check for company indicators
        has_company = any(kw in l_lower for kw in [
            "pvt", "ltd", "inc", "corp", "llc", "technologies", "solutions",
            "software", "systems", "services", "consulting", "company",
        ])
        # Date range pattern (Jan 2022 – Present, 2021-2023, etc.)
        has_date_range = bool(re.search(
            r'(?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec|\d{4})\s*[-–—to]+\s*(?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec|\d{4}|present|current)',
            l_lower,
        ))

        if has_title or has_company or (section_lines and has_date_range):
            if len(line_str) <= 150:
                key = line_str.lower()
                if key not in seen:
                    seen.add(key)
                    exp_list.append({"role_or_company": line_str[:150]})

        if len(exp_list) >= 8:
            break

    return exp_list


# ── Projects Extractor ──

def extract_projects_info(text: str) -> List[dict]:
    exit_names = [
        "skills", "education", "experience", "certifications",
        "hobbies", "declaration", "references", "personal", "achievements",
    ]
    section_lines = _isolate_section(
        text,
        section_names=["projects", "academic projects", "personal projects"],
        exit_names=exit_names,
    )

    if not section_lines:
        return []

    projects: List[dict] = []
    current_title: Optional[str] = None
    current_desc: List[str] = []

    for line_str in section_lines:
        if not line_str or len(line_str) < 3:
            continue

        l_lower = line_str.lower()
        if any(ex in l_lower for ex in ["date of birth", "dob", "email", "phone"]):
            continue

        # Heuristic: a project title is typically short, possibly bold/capitalized
        is_title_like = (
            len(line_str) < 80
            and not line_str.startswith(("•", "-", "●", "○", "▪", "*"))
            and not re.match(r'^\d+\.?\s', line_str)  # not a numbered detail
        )

        # Check if this looks like a new project title (short, no bullet)
        if is_title_like and (
            line_str[0].isupper()
            or re.search(r'(?:project|app|system|platform|tool|website|portal)', l_lower)
        ):
            # Flush previous project
            if current_title:
                desc = " ".join(current_desc).strip()
                detail = f"{current_title}" + (f" — {desc[:120]}" if desc else "")
                projects.append({"details": detail[:200]})

            current_title = line_str
            current_desc = []
        else:
            current_desc.append(line_str)

        if len(projects) >= 6:
            break

    # Flush last project
    if current_title:
        desc = " ".join(current_desc).strip()
        detail = f"{current_title}" + (f" — {desc[:120]}" if desc else "")
        projects.append({"details": detail[:200]})

    return projects[:6]


# ── Certifications Extractor ──

def extract_certifications_info(text: str) -> List[str]:
    # Try section-isolated extraction
    exit_names = [
        "education", "experience", "projects", "skills",
        "hobbies", "declaration", "references", "personal", "summary",
    ]
    section_lines = _isolate_section(
        text,
        section_names=["certifications", "certificates", "licenses"],
        exit_names=exit_names,
    )

    certs: List[str] = []
    seen: set = set()

    # If we found a certifications section, take all non-trivial lines from it
    if section_lines:
        for line_str in section_lines:
            if not line_str or len(line_str) < 5:
                continue
            if EMAIL_REGEX.search(line_str) or PHONE_REGEX.search(line_str):
                continue
            # Clean bullet points
            cleaned = re.sub(r'^[•\-●○▪*]\s*', '', line_str).strip()
            if cleaned and cleaned.lower() not in seen:
                seen.add(cleaned.lower())
                certs.append(cleaned[:120])
            if len(certs) >= 8:
                break
    else:
        # Fallback: scan entire document for certification keywords
        cert_keywords = [
            "certified", "certification", "certificate", "aws",
            "azure", "google cloud", "coursera", "udemy", "edx",
            "cisco", "oracle", "scrum", "pmp", "comptia",
            "meta", "ibm", "microsoft certified", "hackerrank",
            "nptel", "swayam",
        ]
        for line in text.splitlines():
            line_str = line.strip()
            if not line_str or len(line_str) < 5:
                continue
            l_lower = line_str.lower()
            if any(kw in l_lower for kw in cert_keywords):
                if len(line_str) <= 120 and not EMAIL_REGEX.search(line_str):
                    key = line_str.lower()
                    if key not in seen:
                        seen.add(key)
                        certs.append(line_str)
            if len(certs) >= 8:
                break

    return certs


# ------------------------------------------------------------------
# Resume Parser + ATS
# ------------------------------------------------------------------

def parse_resume(filename: str, file_bytes: bytes) -> dict:

    # 1. Extract text
    raw_text = extract_text(filename, file_bytes)

    # 2. Clean text for ATS analysis
    cleaned_text = clean_resume_text(raw_text)

    # 3. Comprehensive extraction (95%+ details)
    name = extract_name(raw_text)
    email = extract_email(raw_text)
    phone = extract_phone(raw_text)
    education_entries = extract_education_history(raw_text)
    experience_entries = extract_experience_history(raw_text)
    project_entries = extract_projects_info(raw_text)
    certification_entries = extract_certifications_info(raw_text)

    # 4. ATS Analysis
    contact = detect_contact_info(raw_text)

    sections = detect_sections(cleaned_text)

    # 4. Extract skills with rule-based detector & spaCy NER
    skills = detect_skills(cleaned_text)
    try:
        from app.ai.spacy_parser import parse_resume_with_spacy
        spacy_result = parse_resume_with_spacy(cleaned_text)
        if spacy_result.get("skills"):
            combined = set(skills + spacy_result["skills"])
            skills = sorted(list(combined))
    except Exception as exc:
        logger.debug("spaCy skill extraction skipped (%s), using rule-based detected skills", exc)

    # 5. Scores
    contact_points = contact_score(contact)

    section_points = section_score(sections)

    skill_points = skill_score(skills)

    ats_score = calculate_ats_score(
        contact_points,
        section_points,
        skill_points,
    )

    # 6. Suggestions
    suggestions = generate_suggestions(
        contact,
        sections,
        skills,
    )

    # 7. Return complete extracted info
    return {

        "raw_text": raw_text,

        "extracted_name": name,

        "extracted_email": email,

        "extracted_phone": phone,

        "extracted_skills": skills,

        "extracted_education": education_entries,

        "extracted_experience": experience_entries,

        "extracted_projects": project_entries,

        "extracted_certifications": certification_entries,

        "ats": {

            "score": ats_score,

            "contact": contact,

            "sections": sections,

            "skills": skills,

            "suggestions": suggestions

        }

    }