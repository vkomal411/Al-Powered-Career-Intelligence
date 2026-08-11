"""
Resume Formatter Service.
Compiles database structures and generates DOCX, JSON, and Text exports.
"""

import io
from typing import Dict, Any


class ResumeFormatter:

    def compile_resume_data(self, resume_builder_obj: Any) -> Dict[str, Any]:
        data = {
            "name": "Candidate",
            "email": "",
            "phone": "",
            "location": "",
            "target_role": getattr(resume_builder_obj, "target_role", "Software Engineer"),
            "summary": "",
            "experience": [],
            "skills": [],
            "education": []
        }

        if hasattr(resume_builder_obj, "sections"):
            for sec in resume_builder_obj.sections:
                sec_type = getattr(sec, "section_type", "").lower()
                content = getattr(sec, "content", {}) or {}

                if sec_type == "summary":
                    data["summary"] = content.get("text", "") or content.get("summary", "")
                    data["name"] = content.get("name", "") or data["name"]
                    data["email"] = content.get("email", "")
                    data["phone"] = content.get("phone", "")
                    data["location"] = content.get("location", "")

                elif sec_type == "experience":
                    exp_entries = getattr(sec, "experience_entries", [])
                    if exp_entries:
                        data["experience"] = [
                            {
                                "job_title": getattr(exp, "job_title", ""),
                                "company": getattr(exp, "company", ""),
                                "location": getattr(exp, "location", ""),
                                "start_date": getattr(exp, "start_date", ""),
                                "end_date": getattr(exp, "end_date", ""),
                                "description": getattr(exp, "description", ""),
                                "bullets": getattr(exp, "bullets", []) or []
                            }
                            for exp in exp_entries
                        ]

                elif sec_type == "skills":
                    skill_entries = getattr(sec, "skill_entries", [])
                    if skill_entries:
                        data["skills"] = [getattr(s, "skill_name", "") for s in skill_entries if getattr(s, "skill_name", "")]

        return data

    def create_docx(self, data: Dict[str, Any]) -> bytes:
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor

            doc = Document()

            # Name Header
            h_para = doc.add_paragraph()
            h_run = h_para.add_run(data.get("name", "Candidate"))
            h_run.font.size = Pt(18)
            h_run.font.bold = True
            h_run.font.color.rgb = RGBColor(30, 27, 75)

            # Target Role Subtitle
            r_para = doc.add_paragraph()
            r_run = r_para.add_run(data.get("target_role", "Software Engineer"))
            r_run.font.size = Pt(12)
            r_run.font.bold = True
            r_run.font.color.rgb = RGBColor(79, 70, 229)

            # Contact Line
            c_para = doc.add_paragraph()
            c_run = c_para.add_run(f"Email: {data.get('email', '')} | Phone: {data.get('phone', '')}")
            c_run.font.size = Pt(10)
            c_run.font.color.rgb = RGBColor(100, 116, 139)

            # Summary Section
            if data.get("summary"):
                doc.add_heading("PROFESSIONAL SUMMARY", level=2)
                doc.add_paragraph(data["summary"])

            # Experience Section
            if data.get("experience"):
                doc.add_heading("WORK EXPERIENCE", level=2)
                for job in data["experience"]:
                    j_para = doc.add_paragraph()
                    j_run = j_para.add_run(f"{job.get('job_title', '')} — {job.get('company', '')}")
                    j_run.bold = True
                    j_para.add_run(f" ({job.get('start_date', '')} - {job.get('end_date', 'Present')})")

                    if job.get("description"):
                        doc.add_paragraph(job["description"])

                    for bullet in job.get("bullets", []):
                        doc.add_paragraph(bullet, style="List Bullet")

            # Skills Section
            if data.get("skills"):
                doc.add_heading("TECHNICAL SKILLS", level=2)
                doc.add_paragraph(" • ".join(data["skills"]))

            buf = io.BytesIO()
            doc.save(buf)
            buf.seek(0)
            return buf.getvalue()

        except Exception:
            # Fallback simple byte generation
            text_content = f"{data.get('name')}\n{data.get('summary')}"
            return text_content.encode("utf-8")

    def create_txt(self, data: Dict[str, Any]) -> str:
        lines = [
            "=" * 50,
            f"{data.get('name', 'Candidate')} — {data.get('target_role', '')}",
            f"Email: {data.get('email', '')} | Phone: {data.get('phone', '')}",
            "=" * 50,
            "",
            "SUMMARY",
            data.get("summary", ""),
            "",
            "SKILLS",
            " • ".join(data.get("skills", [])),
            "",
            "EXPERIENCE"
        ]
        for job in data.get("experience", []):
            lines.append(f"{job.get('job_title')} at {job.get('company')} ({job.get('start_date')} - {job.get('end_date', 'Present')})")
            lines.append(f"- {job.get('description')}")
            for b in job.get("bullets", []):
                lines.append(f"  • {b}")

        return "\n".join(lines)
