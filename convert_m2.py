import os
import markdown2

input_file = r"c:\Users\venka\OneDrive\Documents\career.AI\CareerPilot_AI_Milestone2_Report.md"
output_file_doc = r"c:\Users\venka\OneDrive\Documents\career.AI\CareerPilot_AI_Milestone2_Report.doc"
output_file_docx = r"c:\Users\venka\OneDrive\Documents\career.AI\CareerPilot_AI_Milestone2_Report.docx"

# Read markdown
with open(input_file, 'r', encoding='utf-8') as f:
    md_text = f.read()

# Convert markdown to HTML for Word .doc
html = markdown2.markdown(md_text, extras=["tables", "fenced-code-blocks", "code-friendly"])

css = """
<style>
    body { font-family: 'Segoe UI', Arial, sans-serif; line-height: 1.6; color: #1e293b; margin: 40px; }
    h1 { color: #0f172a; border-bottom: 2px solid #3b82f6; padding-bottom: 8px; }
    h2 { color: #1e3a8a; margin-top: 24px; border-bottom: 1px solid #e2e8f0; padding-bottom: 4px; }
    h3 { color: #2563eb; margin-top: 18px; }
    table { border-collapse: collapse; width: 100%; margin: 16px 0; }
    th, td { border: 1px solid #cbd5e1; padding: 10px 14px; text-align: left; }
    th { background-color: #f1f5f9; font-weight: bold; color: #0f172a; }
    tr:nth-child(even) { background-color: #f8fafc; }
    pre, code { background-color: #f1f5f9; font-family: 'Consolas', 'Courier New', monospace; border-radius: 4px; }
    pre { padding: 14px; overflow-x: auto; border: 1px solid #e2e8f0; }
    code { padding: 2px 6px; color: #0f172a; }
    blockquote { border-left: 4px solid #3b82f6; background-color: #eff6ff; padding: 10px 16px; margin: 16px 0; color: #1e40af; }
    img { max-width: 100%; height: auto; display: block; margin: 16px 0; }
</style>
"""

with open(output_file_doc, 'w', encoding='utf-8') as f:
    f.write('<!DOCTYPE html>\n<html><head><meta charset="utf-8">' + css + '</head><body>\n')
    f.write(html)
    f.write('\n</body></html>')

print(f"Successfully generated Word HTML format document: {output_file_doc}")

# Also try converting to docx using python-docx if available
try:
    import docx
    from docx.shared import Inches, Pt, RGBColor
    
    doc = docx.Document()
    
    # Process lines
    lines = md_text.split('\n')
    in_code = False
    code_block = []
    
    for line in lines:
        if line.startswith('```'):
            if in_code:
                # end code block
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.2)
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(4)
                run = p.add_run('\n'.join(code_block))
                run.font.name = 'Consolas'
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(30, 41, 59)
                code_block = []
                in_code = False
            else:
                in_code = True
                code_block = []
            continue
            
        if in_code:
            code_block.append(line)
            continue
            
        if line.startswith('# '):
            h = doc.add_heading(line[2:], level=1)
            h.style.font.name = 'Segoe UI'
            h.style.font.color.rgb = RGBColor(15, 23, 42)
        elif line.startswith('## '):
            h = doc.add_heading(line[3:], level=2)
            h.style.font.name = 'Segoe UI'
            h.style.font.color.rgb = RGBColor(30, 58, 138)
        elif line.startswith('### '):
            h = doc.add_heading(line[4:], level=3)
            h.style.font.name = 'Segoe UI'
            h.style.font.color.rgb = RGBColor(37, 99, 235)
        elif line.startswith('> '):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            run = p.add_run(line[2:])
            run.italic = True
            run.font.color.rgb = RGBColor(30, 64, 175)
        elif line.strip():
            doc.add_paragraph(line)
            
    doc.save(output_file_docx)
    print(f"Successfully generated native DOCX document: {output_file_docx}")
except Exception as e:
    print(f"Note on python-docx: {e}")
