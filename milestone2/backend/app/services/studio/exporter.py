"""
Studio Exporter Service.
Generates PDF, Word (.docx), TXT, and JSON Resume exports.
"""

import io
import json
from typing import Dict, List, Any


class StudioExporterService:

    def export_pdf(self, resume_data: Dict[str, Any]) -> bytes:
        """Generates a beautifully styled, high-impact PDF resume suitable for ATS parsing and human review."""
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib import colors
            from reportlab.lib.enums import TA_LEFT

            buffer = io.BytesIO()
            doc = SimpleDocTemplate(
                buffer,
                pagesize=letter,
                rightMargin=36,
                leftMargin=36,
                topMargin=36,
                bottomMargin=36
            )

            styles = getSampleStyleSheet()

            name_style = ParagraphStyle(
                'HeaderName',
                parent=styles['Normal'],
                fontName='Helvetica-Bold',
                fontSize=20,
                leading=24,
                textColor=colors.HexColor('#1e1b4b'),
                alignment=TA_LEFT
            )

            role_style = ParagraphStyle(
                'HeaderRole',
                parent=styles['Normal'],
                fontName='Helvetica-Bold',
                fontSize=11,
                leading=14,
                textColor=colors.HexColor('#4f46e5'),
                alignment=TA_LEFT
            )

            contact_style = ParagraphStyle(
                'HeaderContact',
                parent=styles['Normal'],
                fontName='Helvetica',
                fontSize=9,
                leading=12,
                textColor=colors.HexColor('#64748b'),
                alignment=TA_LEFT
            )

            section_heading_style = ParagraphStyle(
                'SectionHeading',
                parent=styles['Normal'],
                fontName='Helvetica-Bold',
                fontSize=10,
                leading=13,
                textColor=colors.HexColor('#1e293b'),
                spaceBefore=10,
                spaceAfter=4
            )

            body_style = ParagraphStyle(
                'BodyText',
                parent=styles['Normal'],
                fontName='Helvetica',
                fontSize=9.5,
                leading=13.5,
                textColor=colors.HexColor('#334155'),
                spaceAfter=4
            )

            item_header_style = ParagraphStyle(
                'ItemHeader',
                parent=styles['Normal'],
                fontName='Helvetica-Bold',
                fontSize=9.5,
                leading=13.5,
                textColor=colors.HexColor('#0f172a')
            )

            skill_chip_style = ParagraphStyle(
                'SkillChip',
                parent=styles['Normal'],
                fontName='Helvetica',
                fontSize=9,
                leading=13,
                textColor=colors.HexColor('#334155')
            )

            story = []

            # 1. Header Section
            full_name = resume_data.get("full_name") or resume_data.get("fullName") or "Candidate Name"
            target_role = resume_data.get("target_role") or resume_data.get("targetRole") or "Software Professional"
            email = resume_data.get("email", "")
            phone = resume_data.get("phone", "")
            github_url = resume_data.get("github_url") or resume_data.get("githubUrl") or ""

            contact_parts = []
            if email:
                contact_parts.append(f"Email: {email}")
            if phone:
                contact_parts.append(f"Phone: {phone}")
            if github_url:
                contact_parts.append(f"GitHub: {github_url}")

            story.append(Paragraph(str(full_name), name_style))
            story.append(Spacer(1, 2))
            story.append(Paragraph(str(target_role), role_style))
            story.append(Spacer(1, 3))
            story.append(Paragraph(" &nbsp;|&nbsp; ".join(contact_parts), contact_style))
            story.append(Spacer(1, 8))
            story.append(HRFlowable(width="100%", thickness=1.5, color=colors.HexColor("#3b82f6"), spaceAfter=10))

            # 2. Professional Summary
            summary = resume_data.get("summary")
            if summary:
                story.append(Paragraph("PROFESSIONAL SUMMARY", section_heading_style))
                story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#cbd5e1"), spaceAfter=6))
                story.append(Paragraph(str(summary), body_style))
                story.append(Spacer(1, 8))

            # 3. Technical Competencies
            skills = resume_data.get("skills", [])
            if skills:
                story.append(Paragraph("TECHNICAL COMPETENCIES", section_heading_style))
                story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#cbd5e1"), spaceAfter=6))
                skill_names = [s["name"] if isinstance(s, dict) else str(s) for s in skills]
                story.append(Paragraph(" &nbsp;|&nbsp; ".join(skill_names), skill_chip_style))
                story.append(Spacer(1, 8))

            # 4. Work Experience
            exps = resume_data.get("experiences") or resume_data.get("experience") or []
            if exps:
                story.append(Paragraph("WORK EXPERIENCE", section_heading_style))
                story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#cbd5e1"), spaceAfter=6))
                for job in exps:
                    title = job.get("job_title") or job.get("jobTitle") or job.get("role") or "Position"
                    comp = job.get("company", "")
                    s_date = job.get("start_date") or job.get("startDate") or ""
                    e_date = job.get("end_date") or job.get("endDate") or "Present"

                    header_text = f"<b>{title}</b> &nbsp;-&nbsp; <i>{comp}</i> <font color='#64748b'>({s_date} - {e_date})</font>"
                    story.append(Paragraph(header_text, item_header_style))
                    if job.get("description"):
                        story.append(Spacer(1, 2))
                        story.append(Paragraph(job["description"], body_style))
                    for bullet in job.get("bullets", []):
                        story.append(Paragraph(f"- {bullet}", body_style))
                    story.append(Spacer(1, 6))

            # 5. Key Projects
            projs = resume_data.get("projects") or resume_data.get("projectsList") or []
            if projs:
                story.append(Paragraph("KEY PROJECTS", section_heading_style))
                story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#cbd5e1"), spaceAfter=6))
                for p in projs:
                    p_name = p.get("name", "Project")
                    techs = p.get("technologies", "")
                    tech_str = f" <font color='#4f46e5'>[{techs}]</font>" if techs else ""
                    story.append(Paragraph(f"<b>{p_name}</b>{tech_str}", item_header_style))
                    if p.get("description"):
                        story.append(Spacer(1, 2))
                        story.append(Paragraph(p["description"], body_style))
                    story.append(Spacer(1, 6))

            # 6. Education
            edu = resume_data.get("education") or resume_data.get("education_list") or resume_data.get("educationList") or []
            if edu:
                story.append(Paragraph("EDUCATION", section_heading_style))
                story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#cbd5e1"), spaceAfter=6))
                for ed in edu:
                    degree = ed.get("degree", "Degree")
                    school = ed.get("school") or ed.get("institution") or ""
                    yr = ed.get("year") or ed.get("graduation_year") or ""
                    yr_str = f" ({yr})" if yr else ""
                    story.append(Paragraph(f"<b>{degree}</b> &nbsp;-&nbsp; {school}{yr_str}", body_style))
                    story.append(Spacer(1, 4))

            doc.build(story)
            buffer.seek(0)
            return buffer.getvalue()

        except Exception:
            return self._export_pdf_fallback(resume_data)

    def _export_pdf_fallback(self, resume_data: Dict[str, Any]) -> bytes:
        """Textwrap-based clean fallback PDF generator ensuring word-wrapped lines and ASCII bullets."""
        import textwrap

        full_name = resume_data.get("full_name") or resume_data.get("fullName") or "Candidate"
        target_role = resume_data.get("target_role") or resume_data.get("targetRole") or "Software Professional"
        email = resume_data.get("email", "")
        phone = resume_data.get("phone", "")

        raw_lines = [
            str(full_name),
            str(target_role),
            f"Email: {email} | Phone: {phone}",
        ]
        if resume_data.get("summary"):
            raw_lines += ["", "PROFESSIONAL SUMMARY", str(resume_data["summary"])]
        if resume_data.get("skills"):
            skills = [s["name"] if isinstance(s, dict) else str(s) for s in resume_data["skills"]]
            raw_lines += ["", "TECHNICAL COMPETENCIES", " | ".join(skills)]

        exps = resume_data.get("experiences") or resume_data.get("experience") or []
        if exps:
            raw_lines += ["", "WORK EXPERIENCE"]
            for item in exps:
                title = item.get("job_title") or item.get("jobTitle") or item.get("role") or "Position"
                comp = item.get("company", "")
                s_date = item.get("start_date") or item.get("startDate") or ""
                e_date = item.get("end_date") or item.get("endDate") or "Present"
                raw_lines.append(f"{title} - {comp} ({s_date} - {e_date})")
                if item.get("description"):
                    raw_lines.append(item["description"])
                raw_lines.extend(f"* {bullet}" for bullet in item.get("bullets", []))

        projs = resume_data.get("projects") or resume_data.get("projectsList") or []
        if projs:
            raw_lines += ["", "KEY PROJECTS"]
            for item in projs:
                p_name = item.get("name", "Project")
                techs = item.get("technologies", "")
                tech_str = f" [{techs}]" if techs else ""
                raw_lines.append(f"{p_name}{tech_str}")
                if item.get("description"):
                    raw_lines.append(item["description"])

        edu = resume_data.get("education") or resume_data.get("education_list") or resume_data.get("educationList") or []
        if edu:
            raw_lines += ["", "EDUCATION"]
            for item in edu:
                degree = item.get("degree", "Degree")
                school = item.get("school") or item.get("institution") or ""
                yr = item.get("year") or item.get("graduation_year") or ""
                yr_str = f" ({yr})" if yr else ""
                raw_lines.append(f"{degree} - {school}{yr_str}")

        wrapped = []
        for line in raw_lines:
            if not line.strip():
                wrapped.append("")
                continue
            for sub in textwrap.wrap(str(line), width=85, replace_whitespace=False):
                wrapped.append(sub)

        pages = [wrapped[i:i + 48] for i in range(0, len(wrapped), 48)] or [[]]

        objects = []
        page_ids = []
        content_ids = []
        for page in pages:
            content = ["BT", "/F1 10 Tf", "50 750 Td", "14 TL"]
            for line in page:
                safe = line.encode("ascii", "replace").decode("ascii").replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
                content.append(f"({safe}) Tj T*" if safe else "T*")
            content.append("ET")
            content_bytes = "\n".join(content).encode("ascii")
            content_ids.append(len(objects) + 1)
            objects.append(f"<< /Length {len(content_bytes)} >>\nstream\n".encode() + content_bytes + b"\nendstream")
            page_ids.append(len(objects) + 1)
            objects.append(b"")

        pages_id = len(objects) + 1
        objects.append(b"")
        font_id = len(objects) + 1
        objects.append(b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>")
        for idx, page_id in enumerate(page_ids):
            objects[page_id - 1] = f"<< /Type /Page /Parent {pages_id} 0 R /MediaBox [0 0 612 792] /Resources << /Font << /F1 {font_id} 0 R >> >> /Contents {content_ids[idx]} 0 R >>".encode()
        kids = " ".join(f"{page_id} 0 R" for page_id in page_ids)
        objects[pages_id - 1] = f"<< /Type /Pages /Kids [{kids}] /Count {len(page_ids)} >>".encode()

        root_id = len(objects) + 1
        objects.append(f"<< /Type /Catalog /Pages {pages_id} 0 R >>".encode())
        pdf = bytearray(b"%PDF-1.4\n")
        offsets = [0]
        for obj_id, obj in enumerate(objects, 1):
            offsets.append(len(pdf))
            pdf += f"{obj_id} 0 obj\n".encode() + obj + b"\nendobj\n"
        xref = len(pdf)
        pdf += f"xref\n0 {len(objects) + 1}\n0000000000 65535 f \n".encode()
        pdf += b"".join(f"{offset:010d} 00000 n \n".encode() for offset in offsets[1:])
        pdf += f"trailer\n<< /Size {len(objects) + 1} /Root {root_id} 0 R >>\nstartxref\n{xref}\n%%EOF".encode()
        return bytes(pdf)

    def export_docx(self, resume_data: Dict[str, Any]) -> bytes:
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor

            doc = Document()
            full_name = resume_data.get("full_name") or resume_data.get("fullName") or "Candidate"
            target_role = resume_data.get("target_role") or resume_data.get("targetRole") or "Software Professional"
            email = resume_data.get("email", "")
            phone = resume_data.get("phone", "")

            # Name Header
            h_para = doc.add_paragraph()
            h_run = h_para.add_run(str(full_name))
            h_run.font.size = Pt(18)
            h_run.font.bold = True
            h_run.font.color.rgb = RGBColor(30, 27, 75)

            # Target Role
            r_para = doc.add_paragraph()
            r_run = r_para.add_run(str(target_role))
            r_run.font.size = Pt(12)
            r_run.font.bold = True
            r_run.font.color.rgb = RGBColor(79, 70, 229)

            # Contact
            c_para = doc.add_paragraph()
            c_run = c_para.add_run(f"Email: {email} | Phone: {phone}")
            c_run.font.size = Pt(10)
            c_run.font.color.rgb = RGBColor(100, 116, 139)

            # Summary
            if resume_data.get("summary"):
                doc.add_heading("PROFESSIONAL SUMMARY", level=2)
                doc.add_paragraph(str(resume_data["summary"]))

            # Skills
            skills = resume_data.get("skills", [])
            if skills:
                doc.add_heading("TECHNICAL COMPETENCIES", level=2)
                skill_names = [s["name"] if isinstance(s, dict) else str(s) for s in skills]
                doc.add_paragraph(" • ".join(skill_names))

            # Experience
            exps = resume_data.get("experiences") or resume_data.get("experience") or []
            if exps:
                doc.add_heading("WORK EXPERIENCE", level=2)
                for job in exps:
                    title = job.get("job_title") or job.get("jobTitle") or job.get("role") or "Position"
                    comp = job.get("company", "")
                    s_date = job.get("start_date") or job.get("startDate") or ""
                    e_date = job.get("end_date") or job.get("endDate") or "Present"

                    j_para = doc.add_paragraph()
                    j_run = j_para.add_run(f"{title} — {comp}")
                    j_run.bold = True
                    j_para.add_run(f" ({s_date} - {e_date})")

                    if job.get("description"):
                        doc.add_paragraph(job["description"])

                    for b in job.get("bullets", []):
                        doc.add_paragraph(b, style="List Bullet")

            # Projects
            projs = resume_data.get("projects") or resume_data.get("projectsList") or []
            if projs:
                doc.add_heading("KEY PROJECTS", level=2)
                for p in projs:
                    p_para = doc.add_paragraph()
                    p_run = p_para.add_run(p.get("name", "Project"))
                    p_run.bold = True
                    if p.get("technologies"):
                        p_para.add_run(f" [{p['technologies']}]")
                    if p.get("description"):
                        doc.add_paragraph(p["description"])

            # Education
            edu = resume_data.get("education") or resume_data.get("education_list") or resume_data.get("educationList") or []
            if edu:
                doc.add_heading("EDUCATION", level=2)
                for ed in edu:
                    degree = ed.get("degree", "Degree")
                    school = ed.get("school") or ed.get("institution") or ""
                    yr = ed.get("year") or ed.get("graduation_year") or ""
                    yr_str = f" ({yr})" if yr else ""
                    doc.add_paragraph(f"{degree} — {school}{yr_str}")

            buf = io.BytesIO()
            doc.save(buf)
            buf.seek(0)
            return buf.getvalue()

        except Exception:
            full_name = resume_data.get("full_name") or resume_data.get("fullName") or "Candidate"
            summary = resume_data.get("summary", "")
            text_str = f"{full_name}\n{summary}"
            return text_str.encode("utf-8")

    def export_txt(self, resume_data: Dict[str, Any]) -> str:
        full_name = resume_data.get("full_name") or resume_data.get("fullName") or "Candidate"
        target_role = resume_data.get("target_role") or resume_data.get("targetRole") or ""
        email = resume_data.get("email", "")
        phone = resume_data.get("phone", "")

        lines = [
            "=" * 50,
            f"{full_name} — {target_role}",
            f"Email: {email} | Phone: {phone}",
            "=" * 50,
            "",
            "SUMMARY",
            resume_data.get("summary", ""),
            "",
            "SKILLS",
            " • ".join([s["name"] if isinstance(s, dict) else str(s) for s in resume_data.get("skills", [])]),
            "",
            "EXPERIENCE"
        ]
        exps = resume_data.get("experiences") or resume_data.get("experience") or []
        for job in exps:
            title = job.get("job_title") or job.get("jobTitle") or job.get("role") or "Position"
            comp = job.get("company", "")
            s_date = job.get("start_date") or job.get("startDate") or ""
            e_date = job.get("end_date") or job.get("endDate") or "Present"
            lines.append(f"{title} at {comp} ({s_date} - {e_date})")
            if job.get("description"):
                lines.append(f"- {job.get('description')}")

        projs = resume_data.get("projects") or resume_data.get("projectsList") or []
        if projs:
            lines.append("")
            lines.append("KEY PROJECTS")
            for p in projs:
                lines.append(f"* {p.get('name')} [{p.get('technologies', '')}]: {p.get('description')}")

        edu = resume_data.get("education") or resume_data.get("education_list") or resume_data.get("educationList") or []
        if edu:
            lines.append("")
            lines.append("EDUCATION")
            for ed in edu:
                degree = ed.get("degree", "Degree")
                school = ed.get("school") or ed.get("institution") or ""
                yr = ed.get("year") or ed.get("graduation_year") or ""
                lines.append(f"* {degree} - {school} ({yr})")

        return "\n".join(lines)

    def export_json(self, resume_data: Dict[str, Any]) -> str:
        return json.dumps(resume_data, indent=2)
